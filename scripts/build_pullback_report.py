#!/usr/bin/env python3
from __future__ import annotations

import csv
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "pullback_20260830"
FIG = ROOT / "analysis_output" / "figures"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "눌림목매매_학습및전략설계_20260830.docx"

NAVY = "17324D"; BLUE = "2E74B5"; DARK_BLUE = "1F4D78"; LIGHT = "EAF1F6"; PALE = "F6F8FA"
GRID = "D9E2EA"; RED = "C94C4C"; GREEN = "4E8D61"; ORANGE = "E69F47"; WHITE = "FFFFFF"; TEXT = "263746"; MUTED = "667788"
FONT = "AppleGothic"

baseline = pd.read_csv(ROOT / "analysis_output/baseline_metrics.csv")
comp = pd.read_csv(ROOT / "analysis_output/strategy_comparison.csv")
bench = pd.read_csv(ROOT / "analysis_output/benchmark_metrics.csv")
walk = pd.read_csv(ROOT / "analysis_output/walk_forward.csv")
robust = pd.read_csv(ROOT / "analysis_output/robustness.csv")
cases = pd.read_csv(ROOT / "analysis_output/example_cases.csv")
sources = pd.read_csv(ROOT / "research_sources.csv")
standalone = pd.read_csv(ROOT / "analysis_output/standalone_filters.csv")
stress = pd.read_csv(ROOT / "analysis_output/cost_stress.csv")
ci = pd.read_csv(ROOT / "analysis_output/expectancy_ci.csv")


def pct(v, d=1):
    if pd.isna(v): return "-"
    return f"{v*100:.{d}f}%"


def num(v, d=2):
    if pd.isna(v): return "-"
    return f"{v:.{d}f}"


def set_run(run, size=None, bold=None, color=None, italic=None, font_name=FONT):
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None: run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa); table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i])); tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[i] / 1440); set_cell_margins(cell)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); hdr = OxmlElement("w:tblHeader"); hdr.set(qn("w:val"), "true"); trPr.append(hdr)


def border_bottom(paragraph, color=NAVY, size=18):
    pPr = paragraph._p.get_or_add_pPr(); pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None: pBdr = OxmlElement("w:pBdr"); pPr.append(pBdr)
    bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(size)); bottom.set(qn("w:space"), "6"); bottom.set(qn("w:color"), color); pBdr.append(bottom)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part; rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), rid)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), BLUE); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u); r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; r.append(t); hyperlink.append(r); paragraph._p.append(hyperlink)


def page_field(paragraph):
    run = paragraph.add_run(); fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); run._r.addnext(fld)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)
sec.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]; normal.font.name = FONT; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(TEXT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK_BLUE,8,4)]:
    st=styles[name]; st.font.name=FONT; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st._element.rPr.rFonts.set(qn("w:eastAsia"),FONT)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
for name in ["List Bullet","List Number"]:
    st=styles[name]; st.font.name=FONT; st.font.size=Pt(11); st._element.rPr.rFonts.set(qn("w:eastAsia"),FONT); st.paragraph_format.left_indent=Inches(.5); st.paragraph_format.first_line_indent=Inches(-.25); st.paragraph_format.space_after=Pt(8); st.paragraph_format.line_spacing=1.167
for name in ["Caption"]:
    st=styles[name]; st.font.name=FONT; st.font.size=Pt(9); st.font.italic=True; st.font.color.rgb=RGBColor.from_string(MUTED); st._element.rPr.rFonts.set(qn("w:eastAsia"),FONT)

if "Definition" not in styles:
    st=styles.add_style("Definition", WD_STYLE_TYPE.PARAGRAPH); st.base_style=styles["Normal"]; st.paragraph_format.left_indent=Inches(.18); st.paragraph_format.space_after=Pt(4)
if "Source Citation" not in styles:
    st=styles.add_style("Source Citation", WD_STYLE_TYPE.PARAGRAPH); st.base_style=styles["Normal"]; st.font.size=Pt(9); st.font.color.rgb=RGBColor.from_string(MUTED); st.paragraph_format.space_before=Pt(4); st.paragraph_format.space_after=Pt(4)

header = sec.header
p=header.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
r=p.add_run("PULLBACK RESEARCH"); set_run(r,9,True,NAVY)
r=p.add_run("    |    학습·전략 설계"); set_run(r,9,False,MUTED)
border_bottom(p, GRID, 6)
footer=sec.footer; p=footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=p.add_run("2026-08-30    ·    "); set_run(r,8,False,MUTED); page_field(p)


SUPPRESS_OUTPUT = False


def detached_paragraph():
    return Document().add_paragraph()


def para(text="", bold_prefix=None, italic=False, align=None, color=None, size=None):
    if SUPPRESS_OUTPUT:
        return detached_paragraph()
    p=doc.add_paragraph();
    if align is not None: p.alignment=align
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); set_run(r,size or 11,True,color or TEXT)
        r=p.add_run(text[len(bold_prefix):]); set_run(r,size or 11,False,color or TEXT,italic)
    else:
        r=p.add_run(text); set_run(r,size or 11,None,color or TEXT,italic)
    return p


def callout(label, text, color=BLUE):
    if SUPPRESS_OUTPUT:
        return detached_paragraph()
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.right_indent=Inches(.08); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(10)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); left=OxmlElement("w:left"); left.set(qn("w:val"),"single"); left.set(qn("w:sz"),"28"); left.set(qn("w:space"),"8"); left.set(qn("w:color"),color); pBdr.append(left); pPr.append(pBdr)
    r=p.add_run(label+"  "); set_run(r,11,True,color); r=p.add_run(text); set_run(r,11,False,TEXT)
    return p


def term(ko, en, abbr, explanation):
    if SUPPRESS_OUTPUT:
        return detached_paragraph()
    p=doc.add_paragraph(style="Definition")
    head=f"{ko}({en}{', '+abbr if abbr else ''}): "
    r=p.add_run(head); set_run(r,10.5,True,DARK_BLUE); r=p.add_run(explanation); set_run(r,10.5,False,TEXT)


def bullet(text):
    if SUPPRESS_OUTPUT:
        return detached_paragraph()
    p=doc.add_paragraph(style="List Bullet"); r=p.add_run(text); set_run(r,10.5,False,TEXT); return p


def numbered(text):
    if SUPPRESS_OUTPUT:
        return detached_paragraph()
    p=doc.add_paragraph(style="List Number"); r=p.add_run(text); set_run(r,10.5,False,TEXT); return p


def heading(text, level=1, page_break=False):
    if SUPPRESS_OUTPUT:
        return detached_paragraph()
    if page_break: doc.add_page_break()
    p=doc.add_paragraph(text, style=f"Heading {level}")
    return p


def table(headers, rows, widths=None, font_size=8.7, header_fill=LIGHT):
    if SUPPRESS_OUTPUT:
        return Document().add_table(rows=1, cols=max(1, len(headers)))
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); shade(c,header_fill); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: set_run(r,font_size,True,NAVY)
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text="" if v is None else str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: set_run(r,font_size,False,TEXT)
    if widths is None:
        base=9360//len(headers); widths=[base]*len(headers); widths[-1]+=9360-sum(widths)
    set_table_geometry(t,widths)
    return t


def figure(filename, caption, width=6.5, page_break=False):
    if SUPPRESS_OUTPUT:
        return None
    if page_break: doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    run=p.add_run(); inline=run.add_picture(str(FIG/filename),width=Inches(width))
    try:
        docPr=inline._inline.docPr; docPr.set("descr",caption)
    except Exception: pass
    cp=doc.add_paragraph(caption,style="Caption"); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.keep_with_next=False


def section_intro(terms, conclusion):
    heading("핵심 용어",2)
    for x in terms: term(*x)
    callout("한 줄 결론", conclusion)


def meaning_and_caution(meaning, caution):
    heading("실제 매매에서의 의미",2); para(meaning)
    heading("해석상의 주의점",2); para(caution)


# Cover: editorial_cover pattern + standard_business_brief preset
for _ in range(5): para("")
p=para("QUANT RESEARCH · FIELD GUIDE",align=WD_ALIGN_PARAGRAPH.CENTER,color=ORANGE,size=11); p.paragraph_format.space_after=Pt(18)
p=para("눌림목 매매",align=WD_ALIGN_PARAGRAPH.CENTER,color=NAVY,size=30); p.runs[0].bold=True; p.paragraph_format.space_after=Pt(4)
p=para("학습 및 전략 설계 리서치",align=WD_ALIGN_PARAGRAPH.CENTER,color=BLUE,size=17); p.paragraph_format.space_after=Pt(22)
p=para("가격 패턴 → 참여자 행동 → 객관적 조건 → 실패조건 → 손익구조 → 실행규칙",align=WD_ALIGN_PARAGRAPH.CENTER,color=MUTED,size=10.5); p.paragraph_format.space_after=Pt(70)
p=para("2026년 8월 30일",align=WD_ALIGN_PARAGRAPH.CENTER,color=NAVY,size=12); p.runs[0].bold=True
para("일반 투자자를 위한 교육·연구용 보고서",align=WD_ALIGN_PARAGRAPH.CENTER,color=MUTED,size=10)
for _ in range(3): para("")
callout("중요", "이 보고서는 투자 권유가 아니다. 눌림목의 개념·근거·실행 규칙을 정리한 교육용 자료이며, 특정 시장이나 종목에서의 수익성을 보장하지 않는다.", RED)
doc.add_page_break()

# Reader guide / TOC
heading("이 보고서를 읽는 법",1)
para("눌림목은 미래에 상승한 차트를 골라 붙이는 이름이 아니다. 이 보고서는 후보 상태인 Setup과 실제 매수 신호인 Trigger를 분리하고, 모든 조건을 당시에 계산할 수 있는 값으로 바꾼다.")
table(["구분","이 보고서의 처리"],[
    ["목적","눌림목의 구조·발생 배경·유형·진입·손절·청산을 하나의 실행 체계로 정리"],
    ["근거","모멘텀·52주 신고가·거래량·지지저항·기술규칙 관련 학술연구를 구성요소별 검토"],
    ["원칙","Setup과 Trigger를 분리하고, 모든 조건과 무효화 가격을 주문 전에 기록"],
    ["주의","문헌의 시장·기간·보유기간이 다르므로 개별 규칙의 수익성을 직접 보장하지 않음"],
    ["실행","시장→추세→조정→지지→Trigger→Risk→Reward 순서의 의사결정나무 제공"],
    ["부속자료","전체 출처의 표본·정의·결과·한계·URL을 정리한 출처 색인 제공"],
],[1700,7660],9.5)
heading("목차",1)
toc=["0. Executive Summary","1. 눌림목 매매란 무엇인가","2. 왜 눌림목이 발생하는가","3. 눌림목의 유형","4. 가격 구조","5. 거래량","6. 이동평균선과 추세","7. 상대강도와 시장환경","8. 변동성","9. 진입 방법","10. 손절","11. 익절과 청산","12. 실패하는 눌림목","13. 기존 학술연구","14. 전략 설계 원칙","15. 실전 규칙 예시","16. 최종 체크리스트"]
table(["전반부: 개념·구조","후반부: 실행·관리"],[[toc[i] if i<len(toc) else "",toc[i+9] if i+9<len(toc) else ""] for i in range(9)],[4680,4680],9.5)

# 0 Executive Summary
heading("0. Executive Summary",1,page_break=True)
section_intro([
    ("설정","Setup","","매수 후보를 정의하는 사전 조건 묶음"),
    ("실행 신호","Trigger","","후보를 실제 주문으로 바꾸는 확인 조건"),
    ("무효화 가격","Invalidation Level","","전략 가정이 틀렸다고 인정하고 청산할 가격"),
],"눌림목은 독립된 만능 패턴이 아니라 ‘사전 추세·통제된 조정·재개 확인·명확한 무효화’를 묶은 조건부 의사결정 체계다.")
heading("가장 중요한 결론",2)
for x in [
    "좋은 눌림은 먼저 상승추세가 존재하고, 조정이 그 추세를 훼손하지 않으며, 매수세 재유입을 확인할 수 있어야 한다.",
    "거래량 감소·이동평균 지지·신고가 근접은 보조 단서다. 어느 하나도 단독으로 미래 상승을 보장하지 않는다.",
    "Setup과 Trigger를 분리하면 지지선 선취매와 추세 재개 확인 매수를 구분할 수 있다.",
    "손절가는 변동성에 맞추되 구조적 저점 아래에 두고, 진입 전에 수량과 최대 손실을 계산한다.",
    "학술연구는 모멘텀과 일부 가격 기준점의 예측력을 지지하지만, 완전한 눌림목 규칙의 거래 가능성을 직접 증명하지는 않는다.",
]: bullet(x)
callout("의사결정", "눌림목이라는 이름보다 사전 추세, 조정 깊이와 기간, 객관적 지지, Trigger, 손절·수량을 한 장의 체크리스트로 기록하는 것이 중요하다.", RED)
figure("01_pullback_structure.png","그림 1. Setup과 Trigger를 분리한 눌림목의 기본 구조")

# Legacy analysis variables are retained only so the omitted backtest-building
# block remains executable; they are not written to the revised report.
us=baseline[(baseline.market=="US")&(baseline.period=="OOS")].iloc[0]; kr=baseline[(baseline.market=="KR")&(baseline.period=="OOS")].iloc[0]
usb=bench[(bench.market=="US")&(bench.period=="OOS")].iloc[0]; krb=bench[(bench.market=="KR")&(bench.period=="OOS")].iloc[0]

# 1 Definition
heading("1. 눌림목 매매란 무엇인가",1,page_break=True)
section_intro([
    ("사전 추세","Prior Trend","","눌림이 시작되기 전에 이미 존재한 상승 방향"),
    ("조정","Pullback","","상승추세 안에서 일정 깊이와 기간 동안 가격이 후퇴하는 구간"),
    ("추세 재개","Resumption","","조정 뒤 기존 상승 방향이 다시 나타나는 현상"),
    ("설정","Setup","","좋은 후보를 찾는 조건 묶음"),
    ("실행 신호","Trigger","","실제로 매수하도록 만드는 확인 조건"),
],"‘상승 후 하락’은 눌림목 후보일 뿐이며, 추세 재개가 확인되기 전까지는 성공한 눌림목이라고 부를 수 없다.")
heading("개념과 근거",2)
para("객관적 정의는 세 요소로 나뉜다. 첫째, 고점 이전 20거래일 저점에서 고점까지 15% 이상 상승한다. 둘째, 고점 후 2~10거래일 동안 저가 기준 5~15% 조정하되 장기 상승 방향을 훼손하지 않는다. 셋째, 전일 고가 돌파 같은 Trigger가 종가에 확인되면 다음 날 시가로 진입한다.")
para("이 정의는 미래에 오른 사례만 골라내지 않는다. Setup이 생겼지만 Trigger가 없거나 손절로 끝난 사례도 같은 모집단에 포함한다.")
heading("숫자의 의미",2)
table(["요소","기본 측정식","기본 범위"],[
    ["직전 상승","고점 ÷ 고점 전 20일 최저가 - 1","15% 이상"],["조정 깊이","(고점 - 고점 후 최저가) ÷ 고점","5~15%"],["조정 기간","고점일부터 Setup일까지 거래일 수","2~10일"],["Trigger","당일 종가 > 전일 고가","다음 날 시가 진입"],["무효화","구조적 저점 - 0.1×ATR","가격 이탈 시 청산"],
],[1700,4700,2960],9.3)
meaning_and_caution("차트를 볼 때 ‘지지할 것 같다’가 아니라, Setup 조건과 Trigger 발생 여부를 별개로 체크한다.","5~15%, 2~10일은 사전 등록한 연구 시작점이지 자연법칙이 아니다. 3~10%, 8~15% 등 인접 범위에서도 성과가 유지되는지 별도로 본다.")

# 2 Why
heading("2. 왜 눌림목이 발생하는가",1,page_break=True)
section_intro([
    ("차익실현","Profit Taking","","상승 뒤 보유자가 이익을 확정하기 위해 매도하는 행동"),
    ("기준점 효과","Anchoring","","투자자가 이전 고점·매수가격 같은 눈에 띄는 가격을 판단 기준으로 삼는 경향"),
    ("유동성 공급","Liquidity Provision","","대기 주문이 매수·매도 충격을 흡수하며 거래를 성사시키는 기능"),
],"눌림은 차익실현과 신규 수요의 균형 과정으로 설명할 수 있지만, 그 설명만으로 미래 수익이 보장되지는 않는다.")
heading("개념과 근거",2)
para("상승 뒤 기존 보유자의 차익실현이 나오면 가격이 후퇴한다. 동시에 돌파를 놓친 참여자와 추세 추종자는 이전 저항·이동평균·거래 집중가격 부근에 매수 주문을 놓을 수 있다. 주문장 연구는 객관적 지지·저항 가격과 대기 주문 깊이의 관련성을 보여주지만, 이후 초과수익을 직접 입증하지는 않는다. [S13]")
para("52주 신고가 연구는 눈에 띄는 가격 기준점이 투자자의 과소반응과 연결될 수 있음을 보여준다. 그러나 한국을 포함한 국제 연구에서는 비용 후 유의성이 크게 줄고 한국 성과는 유의하지 않았다. [S06][S07]")
meaning_and_caution("행동 설명은 Setup의 논리적 가설을 만든다. 실제 주문은 Trigger와 손절이 확인될 때만 낸다.","‘매물 소화’나 ‘세력의 의도’처럼 측정할 수 없는 서사는 실행 조건이 아니다. 관찰 가능한 가격·거래량으로만 대체한다.")

# 3 Types
heading("3. 눌림목의 유형",1,page_break=True)
section_intro([
    ("되돌림 확인","Retest","","돌파한 가격대로 다시 내려와 지지 여부를 시험하는 과정"),
    ("변동성 수축","Volatility Contraction","","조정이 진행되며 일중 가격 범위가 좁아지는 현상"),
    ("박스권","Trading Range","","일정 상단과 하단 사이에서 가격이 반복되는 횡보 구간"),
],"유형은 모양이 아니라 사전 조건·지지 기준·실패 신호로 구분해야 하며, 급등 후 깊은 조정은 정상 눌림과 별도 취급한다.")
figure("02_pullback_types.png","그림 2. 눌림목 8개 유형의 개념 도식")
table(["유형","사전 조건","조정·거래량","진입 후보","실패 신호"],[
    ["A 이동평균","이동평균 상승","MA 부근; 거래량은 별도 검증","MA 재돌파/전일 고가","MA 하향·저점 이탈"],
    ["B 돌파 후 되돌림","저항 돌파","돌파가 부근 retest","저항 재회복","저항 아래 종가 지속"],
    ["C 얕은 조정","강한 직전 상승","5~8% 짧은 조정","단기 저항 돌파","고변동 대량매도"],
    ["D 거래량 감소","상승구간 거래량 확보","비율 ≤0.75 후보","가격 Trigger 필수","감소 없이 하락 확대"],
],[1550,1950,2000,1900,1960],8.4)
table(["유형","사전 조건","조정·거래량","진입 후보","실패 신호"],[
    ["E 변동성 수축","상승추세","ATR·일중폭 감소","수축 상단 돌파","변동성 확대 하락"],
    ["F 박스 돌파","충분한 횡보기간","상단 retest","박스 상단 회복","박스 내부 복귀"],
    ["G 신고가 부근","52주 신고가 근접","15% 이내 조정","전일/단기저항 돌파","시장 약세·상대강도 악화"],
    ["H 급등 후 깊은 조정","단기 과도한 상승","15~20%+·고거래량","원칙상 제외/별도 연구","펌프앤덤프·추세 종료"],
],[1550,1950,2000,1900,1960],8.4)
meaning_and_caution("HTS에서는 먼저 유형을 이름 붙인 뒤, 그 유형에 맞는 지지 기준과 실패 가격을 적는다.","여러 유형을 한 전략으로 섞으면 전제와 무효화 기준이 달라져 일관된 실행이 어려워진다.")

# 4 Price structure
heading("4. 가격 구조",1,page_break=True)
section_intro([
    ("높아진 고점","Higher High","HH","직전 고점보다 높은 새 고점"),
    ("높아진 저점","Higher Low","HL","직전 저점보다 높은 새 저점"),
    ("전환 저점","Pivot Low","","좌우 봉을 확인한 뒤에야 확정되는 국소 저점"),
    ("갭","Gap","","전일 가격 범위와 겹치지 않는 가격 공백"),
],"좋은 가격 구조는 이전 고점·조정 저점·Trigger·손절 가격이 서로 충돌하지 않고 명확하게 숫자로 적히는 구조다.")
heading("개념과 근거",2)
para("이전 저항이 새로운 지지가 되는 현상은 주문 집중과 기준점 효과로 설명할 수 있다. 그러나 직접 연구는 주로 외환 또는 짧은 주문장 표본이며, ‘저항→지지 전환’의 비용 후 초과수익을 충분히 입증하지 않았다. [S12][S13]")
table(["가격 요소","객관화 방법","주의"],[
    ["이전 고점","신호일까지의 20일 최고가","미래 최고가로 다시 그리지 않음"],["박스 상단","20~60일 반복 고점의 허용오차 범위","주관적 선 대신 반복 횟수 필요"],["Pivot Low","좌우 2봉 확인 뒤 인식","최저점 당일에는 알 수 없음"],["갭 유지","조정 저가가 갭 상단 이상","부분·완전 fill을 구분"],["가격대별 거래량","일중 체결 데이터로 구축","일봉 OHLCV만으로 정확히 재현 불가"],
],[1900,3800,3660],9.0)
meaning_and_caution("지지선은 ‘선을 그은 이유’와 ‘이탈 시 행동’을 함께 적을 때만 거래 규칙이 된다.","Volume Profile과 회전율은 이번 데이터로 검증하지 못했다. 관찰 기법을 통계적 edge로 표현하지 않는다.")

# 5 Volume
heading("5. 거래량",1,page_break=True)
section_intro([
    ("거래량 비율","Volume Ratio","","조정구간 평균 거래량 ÷ 직전 상승구간 평균 거래량"),
    ("거래회전율","Turnover","","거래량을 유통 가능한 주식 수와 비교한 비율"),
    ("비정상 거래량","Abnormal Volume","","종목 자체의 평소 수준과 비교해 이례적으로 큰 거래량"),
],"거래량 감소는 매도 압력 약화의 단서가 될 수 있지만, 가격 구조와 추세 재개 신호 없이 단독 매수 근거로 쓰지 않는다.")
heading("숫자의 의미",2)
table(["관찰","가능한 해석","반드시 함께 볼 것"],[
    ["조정 중 거래량 감소","매도 압력 약화 또는 관심 감소","저점 구조, 스프레드, Trigger"],
    ["하락일 대량거래","분산 매도 또는 이벤트 반영","종가 위치, 공시, 다음 날 회복 여부"],
    ["반등일 거래량 증가","신규 수요 유입 가능성","저항 돌파의 종가 유지"],
    ["낮은 절대 거래량","신호가 아니라 체결 위험","평균거래대금, 호가 공백"],
],[2200,3700,3460],9.2)
para("학술문헌의 low volume은 대개 수개월 평균 회전율 또는 비정상 거래량이다. 이것은 조정구간/상승구간 거래량 비율과 다른 변수다. [S09][S10][S11]")
meaning_and_caution("거래량 수축은 보조 필터로만 쓰고, 반등일 거래량 증가는 별도 Trigger 후보로 구분한다.","유통주식수 이력이 없으면 회전율과 매물 교체를 정확히 판단하기 어렵다. 거래량 감소가 단순한 관심 소멸일 수도 있다.")

# 6 MA trend
heading("6. 이동평균선과 추세",1,page_break=True)
section_intro([
    ("단순이동평균","Simple Moving Average","SMA","최근 n일 종가의 산술평균"),
    ("이격도","Distance from Moving Average","","현재 가격이 이동평균에서 떨어진 정도"),
    ("이동평균 배열","Moving Average Alignment","","단기·중기·장기 평균의 상대적 순서"),
],"이동평균 ‘접촉’보다 평균의 상승 기울기와 가격이 장기 추세 위에 있는지가 더 중요하며, 접촉 자체의 공개 edge는 약하다.")
heading("개념과 근거",2)
para("장기 이동평균은 추세 방향, 중기 이동평균은 조정의 위치, 단기 이동평균은 실행 시점을 설명하는 데 주로 사용한다. 같은 가격을 여러 기간으로 반복 계산하므로 각 선의 역할을 사전에 정해야 한다.")
para("공개 연구에서는 200일 이동평균선 접촉 자체의 예측력이 약했다. 이동평균선을 시각적 지지로 곧바로 해석하지 말고, 기울기·가격 구조·Trigger와 함께 봐야 한다. [S20]")
meaning_and_caution("상승 기울기·장기선 상회·Trigger를 함께 확인하고, 이동평균선만 닿았다는 이유로 매수하지 않는다.","이동평균은 모두 가격을 변형한 값이다. 5·10·20·50·60·120·200일선을 한꺼번에 넣으면 중복 신호와 과최적화가 늘어난다.")

# 7 RS market
heading("7. 상대강도와 시장환경",1,page_break=True)
section_intro([
    ("상대강도","Relative Strength","RS","종목 수익률에서 시장 수익률을 뺀 값 또는 횡단면 순위"),
    ("시장 국면","Market Regime","","시장 지수의 추세·변동성에 따라 구분한 환경"),
    ("시장 폭","Market Breadth","","상승 종목 수 등으로 시장 참여의 넓이를 측정한 값"),
],"상승 종목이라도 시장과 업종이 약하면 역풍이 커진다. 시장·업종·종목의 방향을 위에서 아래로 확인한다.")
heading("적용 순서",2)
table(["층","확인 항목","용도"],[
    ["시장","지수의 장기 추세, 시장 폭, 변동성","신규 매수의 공격성 조절"],
    ["업종","업종지수 추세와 상대강도","구조적 순풍 또는 역풍 확인"],
    ["종목","시장·업종 대비 상대수익, 가격 구조","동시 후보의 우선순위"],
],[1600,4200,3560],9.2)
meaning_and_caution("상대강도는 매수 신호보다 후보 순위에 사용하고, 시장 약세에서는 총 위험예산을 줄인다.","상대강도의 기간과 기준지수가 달라지면 순위도 달라진다. 눈에 맞는 기간만 사후 선택하지 않는다.")

# 8 Volatility
heading("8. 변동성",1,page_break=True)
section_intro([
    ("진정범위","True Range","TR","당일 고저폭과 전일 종가 대비 갭을 함께 반영한 범위"),
    ("평균진정범위","Average True Range","ATR","최근 일정 기간의 진정범위 평균"),
    ("실현변동성","Realized Volatility","","과거 일별 수익률의 표준편차를 연율화한 값"),
],"변동성은 ‘좋고 나쁨’보다 손절폭과 보유수량을 맞추는 위험관리 도구로 사용하는 편이 명확하다.")
heading("개념과 근거",2)
para("고정 5% 손절은 모든 종목의 정상 진동 폭을 같다고 가정한다. ATR 손절은 종목별 최근 움직임에 맞추지만, 급격한 변동성 확대 뒤에는 손절폭이 지나치게 넓어질 수 있다. 이때 포지션 크기를 자동으로 줄여야 한다.")
para("변동성이 낮다는 이유만으로 좋은 눌림은 아니다. 유동성 고갈도 ATR을 낮출 수 있으므로, 가격 구조와 평균거래대금을 함께 확인한다.")
heading("포지션 수량",2)
callout("계산식", "1주당 위험 = 진입가 - 손절가. 허용 손실금액 = 계좌자산 × 거래당 위험률. 보유수량 = 허용 손실금액 ÷ 1주당 위험.")
para("예: 계좌 5,000만원, 거래당 위험 0.5%, 진입 50,000원, 손절 47,500원이면 허용 손실은 25만원, 1주당 위험은 2,500원, 이론 수량은 100주다. 단 한 종목 10% 상한이면 100주가 아니라 5,000,000÷50,000=100주로 같고, 더 비싼 종목은 상한이 먼저 작동한다.")
meaning_and_caution("ATR은 손절 위치를 정하는 유일한 기준보다 구조적 저점 아래 완충 폭과 수량 계산에 사용한다.","실적·증자·임상·규제·거래정지 같은 이벤트 갭은 ATR이 막지 못한다. 이벤트 전 보유 축소가 별도 규칙이어야 한다.")

# 9 Entry
heading("9. 진입 방법",1,page_break=True)
section_intro([
    ("선취매","Anticipatory Entry","","지지 예상 가격 부근에서 추세 재개 전에 진입하는 방식"),
    ("확인매매","Confirmation Entry","","반등·저항 돌파를 확인한 뒤 진입하는 방식"),
    ("단기 저항","Short-term Resistance","","최근 3~5일 고가로 정한 가까운 저항 가격"),
],"선취매는 가격이 유리하고 확인매매는 정보가 더 많다. 어느 방식을 쓰든 무효화 가격과 주문 방식을 먼저 고정한다.")
heading("진입 방식 비교",2)
table(["진입","객관적 조건","장점","주요 위험"],[
    ["선취매","사전 지지 가격의 지정가","가격과 손익비가 유리","추세가 재개되지 않을 수 있음"],
    ["전일 고가 돌파","종가 > 전일 고가","단순하고 재현 가능","가짜 반등"],
    ["단기 저항 돌파","종가 > 최근 3~5일 고가","확인 강도 높음","진입 지연·갭"],
    ["MA 재돌파","종가가 선택 MA 위로 복귀","추세 위치 확인","횡보장에서 반복 신호"],
    ["Higher Low 확인","저점 상승 뒤 반등","구조 확인","확정에 미래 봉 필요"],
],[1700,3000,2200,2460],8.9)
meaning_and_caution("초보자는 하나의 진입 방식만 선택하고 종가 확인과 주문 시점을 일관되게 기록한다.","장중 주문은 갭·틱·스프레드·체결 순서가 결과에 영향을 준다.")

# 10 Stop
heading("10. 손절",1,page_break=True)
section_intro([
    ("구조적 손절","Structural Stop","","눌림 저점 등 가격 구조가 무효화되는 수준의 손절"),
    ("변동성 손절","Volatility Stop","","진입가에서 ATR 배수만큼 떨어진 가격의 손절"),
    ("갭 손실","Gap Loss","","손절가보다 낮은 시가에서 체결되어 예상보다 커진 손실"),
],"손절은 승률을 높이는 도구가 아니라 1회 실패의 계좌 영향을 제한하는 사전 약속이다.")
table(["손절 방식","장점","주의"],[
    ["구조적 저점 아래","매매 가정과 직접 연결","저점이 멀면 수량을 크게 줄여야 함"],
    ["고정 비율","단순하고 계산 쉬움","종목별 변동성 차이를 무시"],
    ["ATR 배수","변동성에 적응","이벤트 갭과 급격한 변동성 확대를 막지 못함"],
    ["이동평균 이탈","추세 추적에 유용","종가 확인 뒤 다음 날 체결 시 갭 노출"],
],[2200,3300,3860],9.0)
meaning_and_caution("구조적 저점 아래에 작은 ATR 완충을 두고, 손절폭이 과도하면 거래를 건너뛴다.","손절을 진입 뒤 넓히지 않는다. 가격제한폭·거래정지에서는 예정 가격 체결이 불가능할 수 있다.")

# 11 Exit
heading("11. 익절과 청산",1,page_break=True)
section_intro([
    ("고정 목표","Fixed Target","","진입가 대비 정해진 수익률에서 청산"),
    ("위험배수","Risk Multiple","R","초기 손절폭을 1R로 두고 수익·손실을 표현하는 단위"),
    ("추적 손절","Trailing Stop","","가격 상승에 따라 손절가를 올리는 청산 규칙"),
    ("시간 청산","Time Exit","","정해진 보유일수에 도달하면 청산하는 규칙"),
],"청산 규칙은 이익의 크기와 보유시간, 되돌림 허용폭을 맞바꾸므로 진입 논리와 일관된 하나의 주 규칙을 정한다.")
table(["청산 방식","적합한 목적","주의"],[
    ["고정 목표","짧은 반등 수확","큰 추세를 일찍 끊을 수 있음"],
    ["R 배수 목표","초기 위험과 보상을 연결","가격 구조와 무관한 목표가 될 수 있음"],
    ["추적 손절","추세 연장 참여","되돌림을 허용해 미실현 이익 감소"],
    ["이동평균 이탈","추세 기반 보유","신호 지연과 갭 위험"],
    ["시간 청산","재개되지 않는 포지션 정리","늦게 출발하는 추세를 놓칠 수 있음"],
],[2100,3300,3960],9.0)
meaning_and_caution("목표·추적·시간 청산 중 하나를 주 규칙으로 정하고, 손절은 항상 별도 무효화 규칙으로 유지한다.","이전 고점처럼 눈에 띄는 가격이 자동으로 좋은 청산가격이 되는 것은 아니다.")

# 12 failures
heading("12. 실패하는 눌림목",1,page_break=True)
section_intro([
    ("분산 매도","Distribution","","고점 부근에서 큰 물량이 시장에 나오며 보유 주체가 바뀌는 과정"),
    ("거짓 돌파","False Breakout","","저항을 넘은 뒤 빠르게 다시 저항 아래로 내려오는 현상"),
    ("과도한 이격","Excessive Extension","","직전 상승이 정상 변동성보다 지나치게 가파른 상태"),
    ("이벤트 위험","Event Risk","","실적·증자·임상·규제 등 가격을 불연속적으로 움직이는 위험"),
],"실패 패턴의 공통점은 ‘조정이 깊다’가 아니라 추세·시장·수급·가격 구조 중 하나 이상이 이미 무효화됐다는 점이다.")
table(["실패 유형","객관적 경고","행동"],[
    ["추세 종료","종가<MA200 또는 MA50 하락","Setup 취소"],["대량 매도","하락일 거래량>20일 중앙값 2배","Trigger 있어도 보류"],["거짓 돌파","돌파 후 1~2일 내 저항 아래 종가","즉시 무효화"],["과도한 이격","직전 상승>3ATR/단기 30%+","깊은조정 H형으로 분리"],["약한 시장","지수<MA200·MA200 하락","신규진입 축소"],["이벤트","실적/증자/임상/규제 일정 임박","보유 축소 또는 미진입"],
],[1800,4300,3260],9.0)
meaning_and_caution("성공 사례보다 실패 구조를 먼저 보고, 손절가격과 갭 위험이 감당 가능한지 확인한다.","일봉 자료만으로는 장중 경로·호가 공백·거래정지·공시 시간을 충분히 파악하기 어렵다.")

# 13 academic
heading("13. 기존 학술연구",1,page_break=True)
section_intro([
    ("횡단면 모멘텀","Cross-sectional Momentum","","상대적으로 오른 자산을 사고 덜 오른 자산을 피하거나 공매도하는 전략"),
    ("시계열 모멘텀","Time-series Momentum","","한 자산의 과거 수익 부호로 그 자산의 방향을 정하는 전략"),
    ("데이터 스누핑","Data Snooping","","많은 규칙을 시험한 뒤 우연히 좋은 규칙을 진짜 우위로 착각하는 문제"),
],"학술연구는 추세·신고가·거래량·지지저항의 일부 예측력을 지지하지만, 완전한 눌림목 규칙의 비용 후 성과를 직접 입증하지 않는다.")
academic=sources[sources.category.str.startswith("학술")]
rows=[]
for r in academic.itertuples(index=False):
    rows.append([r.source_id,f"{r.title} ({r.year})",r.period_market,r.key_result,r.limitation])
table(["ID","연구","기간·시장","주요 결과","눌림 적용 한계"],rows,[650,2500,1900,2200,2110],7.7)
heading("핵심 연결",2)
for x in [
    "모멘텀: 상승추세 후보를 찾는 근거는 있으나, 조정 깊이와 Trigger는 별도 검증 대상이다. [S01][S03]",
    "52주 신고가: 미국에서는 예측력이 있었지만 한국 국제표본은 비유의였고 비용 후 대부분 약화됐다. [S06][S07]",
    "거래량: 거래량 수준과 모멘텀 생애주기의 관계는 있지만, 눌림 거래량 비율을 직접 검증한 것은 아니다. [S09][S10]",
    "지지·저항: 주문 집중과 반등확률 근거는 있으나 비용 후 거래수익이 아니다. [S12][S13]",
    "기술규칙: 오래된 구간의 성과가 최근 기간에서 약해졌고 다중검정 위험이 크다. [S14][S15][S16]",
]: bullet(x)
meaning_and_caution("학술 근거는 전략 구성요소의 가설과 위험요인을 이해하는 데 사용한다.","논문의 시장·기간·보유기간·롱숏 여부가 다르면 수익률 숫자를 개별 매매 규칙에 직접 적용하지 않는다.")

# Backtest review, proprietary tests, robustness tables, and test-derived cases
# are intentionally omitted from the revised report at the user's request.
SUPPRESS_OUTPUT = True

# 14 public backtests
heading("14. 기존 Backtest Review",1,page_break=True)
section_intro([
    ("생존편향","Survivorship Bias","","현재까지 살아남은 종목만 과거에도 존재한 것처럼 사용하는 오류"),
    ("슬리피지","Slippage","","신호가격과 실제 체결가격의 차이"),
    ("지정가 체결 편향","Limit Fill Bias","","과거 저가가 지정가에 닿았다는 이유만으로 실제 체결을 가정하는 오류"),
],"공개 눌림목 백테스트는 평균회귀·지정가 선취매가 많고, fully specified·비용 포함·point-in-time·OOS 종목 전략은 드물다.")
pr=sources[sources.category.isin(["공개백테스트","관찰통계"])]
table(["ID","자료","시장·기간","결과","핵심 한계"],[[r.source_id,f"{r.title} ({r.year})",r.period_market,r.key_result,r.limitation] for r in pr.itertuples(index=False)],[650,2600,1900,2100,2110],8.0)
para("공개 SPY 눌림 전략은 OOS CAR 6.31%, MDD -12.41%였지만 반등 확인 전 선취매이고 슬리피지가 명확하지 않았다. 200일선 접촉 연구는 접촉 자체의 edge가 약하다는 반증을 제공한다. [S19][S20]")
meaning_and_caution("공개 숫자는 아이디어 후보로만 쓰고, 같은 데이터·비용·체결 규칙으로 재현되지 않으면 전략의 증거로 채택하지 않는다.","특히 수천 개 파라미터 중 최고 조합만 제시한 자료는 OOS가 없으면 데이터 스누핑 사례로 분류한다.")

# 15 own backtest
heading("15. 자체 Backtest",1,page_break=True)
section_intro([
    ("인샘플","In-sample","IS","전략 구조와 넓은 파라미터 범위를 개발하는 앞 기간"),
    ("표본외","Out-of-sample","OOS","선택에 사용하지 않은 뒤 기간"),
    ("노출도","Exposure","","평균적으로 계좌가 시장에 투자된 비율"),
    ("이익요인","Profit Factor","PF","총이익을 총손실 절댓값으로 나눈 값"),
],"미국 현행 유동성 표본에서는 양의 거래 기대값이 있었지만 지수 초과수익은 없었고, 한국 표본에서는 비용 후 edge를 확인하지 못했다.")
heading("방법",2)
for x in [
    "Universe: 미국 40개, 한국 40개 현행 유동성 주식. 2015~2025 일별 수정 OHLCV. 상장폐지·과거 구성 이력 미포함.",
    "Setup: 최근 고점 전 20일 저점 대비 상승 ≥15%, 고점 후 2~10일, 저가 기준 조정 5~15%, 종가>MA200·지지 MA, MA50 상승.",
    "Trigger: 전일 고가를 종가로 돌파한 날의 다음 거래일 시가.",
    "Stop: 눌림 저점-0.1ATR. Gap down이면 손절가가 아니라 더 낮은 시가 체결.",
    "Exit: 최대 10거래일 보유, 그 전에 Stop 발생 시 청산. 동일 봉 Stop/목표는 Stop 우선.",
    "Portfolio: 최대 동시 10종목, 거래당 10% 명목비중, 상대강도 우선, 동일 종목 중첩 금지.",
    "Cost: 미국 매수/매도 각 0.05%; 한국 매수 0.115%, 매도 0.265%. 배당 총수익 미포함.",
]: bullet(x)
table(["시장·기간","거래","승률","기대값","PF","CAGR","MDD","Sharpe","노출"],[[
    f"{m} {p}",f"{int(r.trades):,}",pct(r.win_rate),pct(r.expectancy,2),num(r.profit_factor),pct(r.cagr,2),pct(r.max_drawdown,2),num(r.sharpe),pct(r.exposure)
] for m in ["US","KR"] for p in ["IS","OOS"] for _,r in baseline[(baseline.market==m)&(baseline.period==p)].iterrows()],[1250,800,900,900,750,900,900,850,1110],8.0)
figure("12_expectancy.png","그림 12. 미국·한국 기본전략과 필터전략의 OOS 기대값")
figure("13_max_drawdown.png","그림 13. 미국·한국 기본전략과 필터전략의 OOS 최대낙폭")
heading("초과수익 판단",2)
para(f"미국 OOS 전략 CAGR {pct(us.cagr,2)}는 S&P 500 {pct(usb.cagr,2)}보다 {pct(us.cagr-usb.cagr,2)} 낮았다. 한국 전략 CAGR {pct(kr.cagr,2)}는 KOSPI {pct(krb.cagr,2)}보다 {pct(kr.cagr-krb.cagr,2)} 낮았다. 낮은 노출도가 낙폭을 줄였지만 ‘현금 대비 양수’와 ‘지수 대비 초과수익’은 다른 질문이다.")
meaning_and_caution("전략은 지수를 대체하는 핵심 보유보다 제한된 위험예산의 전술적 sleeve로만 평가한다.","현재 종목 표본을 과거에 소급한 생존편향 때문에 결과는 시장 대표 백테스트가 아니라 현존 유동성 표본의 연구 결과다.")

# 16 filters
heading("16. Filter Incremental Test",1,page_break=True)
section_intro([
    ("증분 검정","Incremental Test","","기본전략에 필터를 한 번에 하나씩 추가해 변화량을 보는 실험"),
    ("복잡성 비용","Complexity Cost","","조건 증가로 거래 수·해석 가능성·재현성이 줄어드는 비용"),
],"미국에서는 거래량·시장·상대강도가 개선에 기여했지만, 변동성·신고가 필터는 추가 가치가 거의 없었고 한국에는 전체 조합이 이전되지 않았다.")
figure("15_incremental_filters.png","그림 14. Alpha 필터 증분 테스트")
zf=comp[(comp.period=="OOS")&comp.label.str.startswith("Filter")]
table(["단계","미국 거래/기대값/MDD","한국 거래/기대값/MDD","채택 판단"],[[
    f"F{i} {label}",
    f"{int(zf[(zf.market=='US')&zf.label.str.startswith(f'Filter:{i}:')].iloc[0].trades)} / {pct(zf[(zf.market=='US')&zf.label.str.startswith(f'Filter:{i}:')].iloc[0].expectancy,2)} / {pct(zf[(zf.market=='US')&zf.label.str.startswith(f'Filter:{i}:')].iloc[0].max_drawdown,2)}",
    f"{int(zf[(zf.market=='KR')&zf.label.str.startswith(f'Filter:{i}:')].iloc[0].trades)} / {pct(zf[(zf.market=='KR')&zf.label.str.startswith(f'Filter:{i}:')].iloc[0].expectancy,2)} / {pct(zf[(zf.market=='KR')&zf.label.str.startswith(f'Filter:{i}:')].iloc[0].max_drawdown,2)}",
    decision] for i,label,decision in [(0,"기본","기준"),(1,"+거래량","미국만 개선"),(2,"+시장","미국 개선"),(3,"+상대강도","미국 Advanced 후보"),(4,"+변동성","추가가치 미미"),(5,"+신고가","복잡성 대비 미미")]
],[1500,2850,2850,2160],8.6)
meaning_and_caution("미국 Advanced는 F3까지만 사용한다. 비슷한 성과라면 F4·F5는 버린다. 한국은 필터를 넣어 우위를 주장하지 않는다.","F3 미국 OOS는 69회로 표본이 작다. 거래 수 감소가 위험 축소와 우연한 종목 선택을 동시에 만들 수 있다.")

# 17 robustness
heading("17. Robustness Test",1,page_break=True)
section_intro([
    ("파라미터 견고성","Parameter Robustness","","인접한 조건값에서도 성과 방향이 유지되는 성질"),
    ("워크포워드","Walk-forward Test","","과거 창에서 조건을 고르고 다음 연도에서 반복 검증하는 절차"),
    ("군집 부트스트랩","Cluster Bootstrap","","같은 시기 거래의 상관을 보존해 신뢰구간을 추정하는 방법"),
],"미국 성과는 일부 파라미터에서 유지됐지만 워크포워드 연도별 부호가 섞였고, 한국은 일관성이 부족했다.")
figure("04_depth_performance.png","그림 15. 조정 깊이별 기대값")
figure("05_duration_performance.png","그림 16. 조정 기간별 기대값")
heading("워크포워드 결과",2)
table(["시장","양의 테스트 연도","음의 테스트 연도","중앙 테스트 기대값","판단"],[[
    m,str(int((walk[walk.market==m].test_expectancy>0).sum())),str(int((walk[walk.market==m].test_expectancy<=0).sum())),pct(walk[walk.market==m].test_expectancy.median(),2),judg] for m,judg in [("US","방향성 있으나 불안정"),("KR","재현성 부족")]
],[1200,1500,1500,1800,3360],9.2)
u_ci=ci[(ci.market=="US")&(ci.family=="Baseline")].iloc[0]; k_ci=ci[(ci.market=="KR")&(ci.family=="Baseline")].iloc[0]
para(f"월 단위 군집 부트스트랩 95% 기대값 구간은 미국 {pct(u_ci.ci95_low,2)}~{pct(u_ci.ci95_high,2)}, 한국 {pct(k_ci.ci95_low,2)}~{pct(k_ci.ci95_high,2)}다. 미국 기본전략도 0을 포함하므로, 필터전략의 더 좁은 양의 구간을 그대로 확정적 알파로 해석하지 않는다.")
heading("과최적화 방지 규칙",2)
for x in ["단일 최고 CAGR보다 인접 셀의 중앙값과 부호를 본다.","Entry·Stop·Exit를 동시에 전수조합하지 않는다.","OOS를 본 뒤 규칙을 바꾸면 그 기간은 더 이상 OOS가 아니다.","비용 2배·한 거래일 지연·시장별 분리에서 방향을 재확인한다.","세부 집단 30회 미만은 표본 부족으로 표시한다."]: bullet(x)
meaning_and_caution("현재 결과는 다음 미래 구간에서 전진검증할 후보를 고르는 용도다.","White Reality Check·Hansen SPA 같은 다중검정 보정은 이번 산출물에서 직접 계산하지 않았다. 시험한 조합 수를 감안하면 명목 성과를 보수적으로 봐야 한다.")

# 18 Beginner
heading("18. Beginner Strategy",1,page_break=True)
section_intro([
    ("거래당 위험","Risk per Trade","","손절될 때 계좌에서 잃도록 허용한 비율"),
    ("손익비","Reward-to-Risk Ratio","","예상 이익을 초기 손실위험으로 나눈 값"),
],"초보자 전략은 5개 조건과 하나의 손절·시간청산만 사용하며, 한국에서는 edge가 검증되지 않았으므로 모의·소액 전진검증이 전제다.")
heading("Strategy A — Beginner",2)
table(["단계","규칙","실행"],[
    ["1 시장","시장지수>상승 중인 200일선","아니면 신규매수 보류"],
    ["2 추세","종가>200일선; MA50 상승; 직전 20일 상승≥15%","후보 저장"],
    ["3 눌림","고점 후 2~10일; 저가 조정 5~15%","범위 밖이면 제외"],
    ["4 Trigger","종가가 전일 고가 돌파","다음 날 시가"],
    ["5 Risk","손절=눌림저점-0.1ATR; 손절폭≤10%","계좌 위험 0.25~0.5%"],
    ["6 Exit","10거래일 또는 손절 중 먼저 발생","손절 확대 금지"],
],[1200,4800,3360],9.2)
heading("매수하지 않는 조건",2)
for x in ["실적·증자·임상·규제 등 이진 이벤트가 보유기간 안에 예정됨","20일 중앙 거래대금이 내 주문의 20배 미만","하락일 거래량이 평소 2배 이상이며 저가 마감","손절폭이 10%를 넘거나 목표 상승여력이 2R 미만","Trigger 없이 지지선만 닿음","같은 종목의 기존 포지션이 있음"]: bullet(x)
callout("한국 적용", "기본전략 OOS 기대값이 0에 가까웠으므로, 실제 자금은 거래당 위험 0.25% 이하·동시 3종목 이하에서 최소 100회 전진검증 후 판단한다.", RED)

# 19 Advanced
heading("19. Advanced Strategy",1,page_break=True)
section_intro([
    ("알파 필터","Alpha Filter","","기본 신호의 기대수익 또는 위험조정 성과를 개선하려는 추가 조건"),
    ("상대강도 초과수익","Relative Strength Excess Return","","종목 60일 수익률 - 시장 60일 수익률"),
],"Advanced 전략은 미국 표본에서 개선된 세 필터까지만 사용하며, 한국에는 채택하지 않는다.")
heading("Strategy B — Advanced (미국 연구 후보)",2)
table(["구성","규칙","OOS 근거"],[
    ["Baseline","Beginner의 추세·눌림·전일고가 Trigger·구조적 Stop·10일 Exit",f"325회; 기대값 {pct(us.expectancy,2)}"],
    ["+거래량","조정평균/상승평균 ≤0.75",f"105회; 기대값 {pct(zf[(zf.market=='US')&zf.label.str.startswith('Filter:1:')].iloc[0].expectancy,2)}"],
    ["+시장","S&P500>MA200, MA200 20일 상승",f"80회; 기대값 {pct(zf[(zf.market=='US')&zf.label.str.startswith('Filter:2:')].iloc[0].expectancy,2)}"],
    ["+상대강도","60일 종목수익-시장수익 ≥5%",f"69회; 기대값 {pct(zf[(zf.market=='US')&zf.label.str.startswith('Filter:3:')].iloc[0].expectancy,2)}"],
    ["제외","ATR범위·52주고가 필터","추가 개선 거의 없음"],
],[1500,4700,3160],9.2)
para("포지션은 초기 위험 0.5%, 한 종목 명목 10%, 최대 10종목을 상한으로 한다. 동시 신호는 상대강도 순으로 선택한다. 손절폭이 넓을수록 수량을 줄인다.")
heading("한국용 Advanced 판단",2)
para("한국은 F3 기대값이 -0.30%, F4 -1.15%, F5 -0.92%였다. 따라서 복잡한 필터를 추가한 Advanced 전략을 만들지 않는다. 신고가 단독 필터는 +0.33%였지만 다른 필터와 결합 시 소멸했고, 현재 표본 생존편향과 다중검정 때문에 채택 근거가 부족하다.")
meaning_and_caution("미국도 Advanced는 독립 미래 구간에서 재검증할 연구 후보이며, 지수 핵심 보유를 대체하지 않는다.","필터 성과가 좋은 이유가 실제 구조인지 69회 표본의 우연인지 구분하려면 더 긴 point-in-time 데이터가 필요하다.")

# 20 cases
heading("20. 실전 사례",1,page_break=True)
section_intro([
    ("사후편향","Hindsight Bias","","결과를 알고 과거 판단이 쉬웠다고 느끼는 오류"),
    ("당시 정보 집합","Information Set","","Setup·Trigger 시점까지 시장 참여자가 알 수 있었던 데이터"),
],"성공·실패·미진입 사례를 같은 규칙으로 읽어야 패턴 암기가 아니라 의사결정 훈련이 된다.")
figure("16_examples_success.png","그림 17. 성공한 눌림목 5개",page_break=False)
for ctype in ["성공","실패","진입 보류"]:
    heading(f"{ctype} 사례",2,page_break=(ctype!="성공"))
    if ctype=="실패": figure("17_examples_failure.png","그림 18. 실패한 눌림목 5개")
    if ctype=="진입 보류": figure("18_examples_ambiguous.png","그림 19. Setup은 있으나 Trigger가 없어 진입하지 않은 3개")
    for _,r in cases[cases.case_type==ctype].iterrows():
        heading(f"{r['market']} · {r['name']} · Setup {str(r['setup_date'])[:10]}",3)
        table(["당시 알 수 있었던 정보","Setup","Trigger","Stop","이후 결과","배울 점"],[[r.known_info,r.setup,r.trigger,f"{r['stop']:,.2f}" if pd.notna(r['stop']) else "-",r.outcome,r.lesson]],[1900,1300,1500,1000,1700,1960],8.1)
para("사례의 ‘성공’은 전략 정의에 따른 순수익 결과일 뿐, 기업가치나 장기 투자 매력의 평가가 아니다. 미진입 사례는 이후 가격이 올랐는지로 판단하지 않는다.")

# Resume the learning-and-strategy report with non-backtest implementation guidance.
SUPPRESS_OUTPUT = False

# 14 strategy design principles
heading("14. 전략 설계 원칙",1,page_break=True)
section_intro([
    ("필수 조건","Core Condition","","전략의 논리를 구성하며 빠지면 거래하지 않는 조건"),
    ("보조 필터","Supporting Filter","","후보의 질을 구분하지만 단독 매수 근거가 되지 않는 조건"),
    ("무효화","Invalidation","","매매 가정이 깨졌음을 가격으로 확인하는 기준"),
],"규칙은 후보 선정, 실행 확인, 위험 제한, 청산의 네 층으로 분리하고 각 층의 역할이 겹치지 않게 만든다.")
table(["층","질문","권장 요소","흔한 오류"],[
    ["후보 선정","상승추세가 이미 존재하는가?","상승 중인 장기 이동평균, 직전 고점·저점 상승","단기 급등만 보고 추세로 간주"],
    ["조정 품질","정상 조정인가 추세 훼손인가?","깊이·기간·저점 구조, 하락일 거래량","깊은 하락을 싸다는 이유로 포함"],
    ["실행 확인","매수세가 다시 우세해졌는가?","전일 고가·단기저항·MA 재돌파 종가","지지선 접촉만으로 선취매"],
    ["위험 제한","틀렸을 때 어디서 얼마나 잃는가?","구조적 저점, ATR 완충, 계좌 위험","손절가를 진입 후 아래로 이동"],
    ["청산","가설이 맞거나 틀렸음을 어떻게 종료하는가?","사전 목표·추세 추적·시간 제한 중 하나","수익 중 규칙을 임의 변경"],
],[1500,2500,3000,2360],9.0)
heading("보조 필터의 역할",2)
for x in [
    "거래량 감소는 매도 압력 약화의 후보 신호지만, 가격 구조와 Trigger 없이 단독으로 사용하지 않는다.",
    "시장 방향은 개별 종목 신호의 성공 환경을 제한하는 상위 조건으로 사용한다.",
    "상대강도는 같은 날 여러 후보가 나올 때 우선순위를 정하는 도구로 사용한다.",
    "RSI·MACD·스토캐스틱을 동시에 넣으면 같은 가격 정보를 중복 계산할 수 있으므로 최소화한다.",
]: bullet(x)
meaning_and_caution("각 조건 옆에 ‘후보·확인·위험·청산’ 중 하나의 역할을 적으면 중복 조건을 줄일 수 있다.","아래 숫자 범위는 실행 예시이며 특정 시장에서 검증된 최적값이 아니다.")

# 15 illustrative rules
heading("15. 실전 규칙 예시",1,page_break=True)
section_intro([
    ("거래당 위험","Risk per Trade","","손절 시 계좌에서 잃도록 허용한 최대 비율"),
    ("손익비","Reward-to-Risk Ratio","R","예상 이익을 초기 손실 위험으로 나눈 비율"),
],"초보자는 하나의 Setup·하나의 Trigger·하나의 손절 규칙만 사용하고, 모든 숫자를 주문 전에 기록한다.")
heading("Strategy A — 단순형",2)
table(["단계","예시 규칙","실행"],[
    ["1 시장","시장지수 > 상승 중인 MA200","아니면 신규 매수 보류"],
    ["2 추세","종가 > MA200, MA50 상승, 고점·저점 상승","후보 등록"],
    ["3 눌림","예: 고점 후 2~10일, 깊이 5~15%","범위 밖은 다른 유형으로 분리"],
    ["4 Trigger","종가가 전일 고가 또는 단기저항 돌파","다음 거래일 지정 방식대로 주문"],
    ["5 Stop","눌림 저점 아래 + ATR 완충","손절폭이 과도하면 제외"],
    ["6 Size","계좌 위험 0.25~0.5% 이내","수량 = 허용손실 ÷ 주당 위험"],
    ["7 Exit","2R·다음 저항·시간 제한 중 사전 선택","진입 후 임의 변경 금지"],
],[1200,4800,3360],9.2)
heading("Strategy B — 확인 강화형",2)
table(["추가 조건","용도","채택 전 질문"],[
    ["조정 거래량 감소","매도 압력 약화 확인","상승 구간 대비 어떤 기간과 비율로 측정할 것인가?"],
    ["시장 추세 필터","약세장 신규 매수 제한","어떤 지수와 이동평균을 기준으로 할 것인가?"],
    ["상대강도","동시 후보 우선순위","시장·업종 대비 어느 기간의 초과수익을 쓸 것인가?"],
    ["변동성 수축","손절 거리와 구조 안정성 점검","ATR 감소가 유동성 고갈과 구분되는가?"],
    ["돌파 후 되돌림","이전 저항의 지지 전환 확인","종가 기준 회복과 실패 가격이 명확한가?"],
],[2000,2800,4560],9.2)
callout("주의", "확인 강화형 조건은 연구 가설이다. 조건이 많아질수록 신호가 드물고 사후 설명이 쉬워지므로, 실전 기록에서 하나씩만 추가한다.", RED)

# 16 checklist
heading("16. 최종 체크리스트",1,page_break=True)
section_intro([
    ("의사결정나무","Decision Tree","","예/아니오 질문을 순서대로 따라가 행동을 정하는 도구"),
],"시장→종목추세→조정→거래량→Trigger→Risk→Reward 순으로 확인하고, 하나라도 핵심 실패조건이면 거래하지 않는다.")
heading("1~2분 Decision Tree",2)
steps=[
    "시장지수가 상승 중인 200일선 위인가? 아니오 → 신규매수 보류.",
    "종목이 200일선 위이고 50일선이 상승하는가? 아니오 → 제외.",
    "직전 20일 상승이 15% 이상이고 고점 후 2~10일 조정인가? 아니오 → 제외.",
    "조정 저가 기준 깊이가 5~15%인가? 아니오 → 깊은 조정/다른 전략으로 분리.",
    "객관적 지지와 손절가격이 적을 수 있는가? 아니오 → 제외.",
    "전일 고가 또는 단기저항 돌파가 종가로 확인됐는가? 아니오 → 관찰만.",
    "손절폭≤10%, 예상 여력≥2R, 이벤트 위험 없음인가? 아니오 → 제외.",
    "계좌 위험 0.25~0.5%로 수량을 계산했는가? 예 → 다음 날 시가 주문.",
]
for x in steps: numbered(x)
heading("1페이지 실전 체크리스트",2)
table(["영역","체크","판정 기준"],[
    ["시장","☐","지수>상승 중인 MA200"],["종목 추세","☐","종가>MA200; MA50 상승"],["직전 상승","☐","고점·저점 상승과 충분한 상승 탄력"],["눌림 깊이","☐","사전에 정한 정상 조정 범위"],["눌림 기간","☐","사전에 정한 기간 범위"],["거래량","☐","하락일 대량매도 여부와 조정 구간 감소 여부"],["지지","☐","이전 고점/박스 상단/상승 MA 중 하나를 숫자로 기록"],["Trigger","☐","전일 고가 또는 단기저항 종가 돌파"],["Stop","☐","구조적 저점 아래; 손절폭 과대 시 제외"],["Reward","☐","예상여력≥2R"],["수량","☐","계좌 위험 0.25~0.5%; 종목 상한 사전 설정"],["실패","☐","대량매도·시장약세·이벤트 위험 없음"],["청산","☐","목표/추적/시간 규칙 중 하나; 진입 후 변경 없음"],
],[1700,800,6860],8.8)
callout("최종 실행 규칙", "체크리스트를 모두 통과해도 수익이 보장되지는 않는다. 기록→모의 또는 소액 적용→정기 복기 순서로 사용하고, 규칙 밖 거래를 성과와 섞지 않는다.", RED)

# QC and appendices
heading("부록 A. Quality Control 답변",1,page_break=True)
qc=[
    ("Q1 정의가 객관적인가?","예. 상승폭·깊이·기간·Trigger·Stop을 수치화했다."),
    ("Q2 Setup과 Trigger를 구분했는가?","예. 후보 조건과 실제 실행 조건을 별도 단계로 분리했다."),
    ("Q3 거래량의 역할을 과장하지 않았는가?","예. 매도 압력 약화의 보조 단서로만 사용했다."),
    ("Q4 이동평균선 edge를 구분했는가?","예. 접촉 자체의 공개 반증과 기울기·추세 조건을 구분했다."),
    ("Q5 RSI·MACD를 중복 사용했는가?","아니오. 같은 가격 정보를 반복 계산하는 보조지표는 기본 규칙에서 제외했다."),
    ("Q6 손절이 구조와 연결되는가?","예. 눌림 저점 아래를 무효화 기준으로 하고 ATR은 완충에만 사용했다."),
    ("Q7 수량을 진입 전에 정하는가?","예. 계좌 허용손실을 주당 위험으로 나누는 방식을 제시했다."),
    ("Q8 이벤트와 갭 위험을 확인했는가?","예. 실적·증자·임상·규제와 유동성 위험을 매수 제외 조건에 포함했다."),
    ("Q9 문헌의 적용 한계를 밝혔는가?","예. 시장·기간·보유기간 차이와 직접 눌림목 증거의 공백을 명시했다."),
    ("Q10 규칙이 단순한가?","예. 후보·확인·위험·청산의 역할별 최소 조건만 남겼다."),
]
table(["점검 질문","답변"],qc,[3600,5760],9.0)

heading("부록 B. 확인하지 못한 항목과 후속 자료",1,page_break=True)
table(["항목","이번 처리","필요 데이터/후속"],[
    ["거래회전율","개념만 설명","유통주식수의 시점별 이력"],["업종 상대강도","개념만 설명","업종 분류와 업종지수"],["Volume Profile","개념만 설명","일중 체결가격·거래량"],["이벤트 위험","매수 제외 조건으로 제안","공시·실적·임상·증자 캘린더"],["RSI/MACD/ADX 등","최소지표 원칙으로 제외","사용 목적과 중복 여부를 사전 정의"],["호가·체결 용량","유동성 점검 원칙만 설명","호가 깊이·스프레드·평균거래대금"],["시장별 적합성","보편 법칙으로 주장하지 않음","시장 구조·세금·거래 관행별 별도 검토"],["기업가치·촉매","기술적 Setup과 분리","실적·밸류에이션·산업 촉매 분석"],
],[2200,3000,4160],9.0)

heading("부록 C. 주요 출처와 부속 파일",1,page_break=True)
report_sources = sources[sources.category.str.startswith("학술")]
para("자체·공개 백테스트 자료를 제외한 학술 출처 19개의 제목·기간·시장·정의·결과·한계·URL은 Pullback_Academic_Source_Index.xlsx에 수록했다. 아래는 직접 연결한 원문이다.")
for _,r in report_sources.iterrows():
    p=doc.add_paragraph(style="Source Citation")
    rr=p.add_run(f"[{r.source_id}] {r.title} ({r.year}). "); set_run(rr,8.8,True,NAVY)
    add_hyperlink(p,"원문",r.url)
    rr=p.add_run(f" — {r.limitation}"); set_run(rr,8.8,False,MUTED)
heading("부속 파일",2)
table(["파일","역할"],[
    ["Pullback_Academic_Source_Index.xlsx","학술 근거·표본·정의·결과·한계·원문 URL 색인"],
],[3300,6060],9.2)

# Core properties and save
doc.core_properties.title="눌림목 매매 학습 및 전략 설계 리서치"
doc.core_properties.subject="눌림목의 개념, 근거, 전략 설계와 실행 체크리스트"
doc.core_properties.author="Quant Research"
doc.core_properties.keywords="눌림목, pullback, momentum, technical analysis, risk management"
doc.save(DOCX)
print(DOCX)
